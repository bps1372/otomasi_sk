import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import requests

st.set_page_config(page_title="AutoSura1372", layout="centered")

st.title("📝 Auto-Sura (Automatis Surat)")
st.write("Aplikasi ini akan mengotomasi pengisian dokumen SK Kuasa Pengguna Anggaran BPS Kota Solok.")

# --- KONFIGURASI GITHUB ---
GITHUB_RAW_URL = "https://raw.githubusercontent.com/bps1372/otomasi_sk/main/TemplateDokumen.docx"

# 1. Input Data Umum
st.subheader("Informasi Umum Dokumen")
col1, col2 = st.columns(2)

with col1:
    nomor = st.text_input("A.Nomor Dokumen", placeholder="Contoh: 042.1 TAHUN 2026")
    tanggal = st.text_input("B.Tanggal Ditetapkan", placeholder="Contoh: 15 Januari 2026")
    petugas = st.text_input("C.Menetapkan...", placeholder="Contoh: Petugas Sensus Ekonomi 2026")

with col2:
    tentang = st.text_input("D.Judul [tulis dengan huruf besar]", placeholder="Contoh: PETUGAS SENSUS EKONOMI 2026")
    pelaksanaan = st.text_input("E.Pelaksanaan Kegiatan", placeholder="Contoh: Sensus Ekonomi 2026")

# 2. Input Data Tabel (Dinamis dengan Fitur Tambah Baris)
st.subheader("F. Lampiran")

# Inisialisasi Session State
if "df_lampiran" not in st.session_state:
    st.session_state.df_lampiran = pd.DataFrame({
        "Nama/Jabatan": ["si ABCD"],
        "NIP/Golongan": ["199102192019031001 Gol: III/b"],
        "Posisi": ["PML"],
        "Honor": ["50.000"]
    })

# State baru untuk menyimpan data hasil editan agar tidak me-refresh tabel secara terus-menerus
if "edited_df" not in st.session_state:
    st.session_state.edited_df = st.session_state.df_lampiran

col_n1, col_n2 = st.columns([1, 2])
with col_n1:
    baris_baru = st.number_input("Ingin tambah berapa baris kosong?", min_value=1, max_value=500, value=1)
with col_n2:
    st.write("##")
    if st.button("Tambahkan Baris Kosong"):
        new_data = pd.DataFrame({
            "Nama/Jabatan": ["si C"] * baris_baru,
            "NIP/Golongan": ["199102192019031001 Gol: III/b"] * baris_baru,
            "Posisi": ["PML"] * baris_baru,
            "Honor": ["50.000"] * baris_baru
        })
        # Gabungkan data BARU dengan data TERAKHIR yang sudah diedit (bukan df awal)
        st.session_state.df_lampiran = pd.concat([st.session_state.edited_df, new_data], ignore_index=True)
        st.rerun()

# Editor Tabel dengan num_rows="dynamic"
edited_df = st.data_editor(
    st.session_state.df_lampiran, 
    num_rows="dynamic", 
    use_container_width=True, 
    key="editor_key"
)

# Kunci data terbaru ke session state "edited_df" (Bukan df_lampiran agar tabel tidak kehilangan fokus)
st.session_state.edited_df = edited_df

# --- STATE UNTUK DOKUMEN ---
if "doc_ready" not in st.session_state:
    st.session_state.doc_ready = False
if "doc_data" not in st.session_state:
    st.session_state.doc_data = None
if "doc_name" not in st.session_state:
    st.session_state.doc_name = ""

# --- FUNGSI CUSTOM FONT & XML ---
def apply_bookman_font(run, size=11):
    run.font.name = 'Bookman Old Style'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Bookman Old Style')
    run.font.size = Pt(size)

def set_cell_text_with_font(cell, text, keep_next=False):
    cell.text = text
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        if keep_next:
            p.paragraph_format.keep_with_next = True
        for run in p.runs:
            apply_bookman_font(run, size=12)

def set_repeat_table_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)

# 3. Tombol Eksekusi
if st.button("Proses & Buat Dokumen", type="primary"):
    if edited_df.empty:
        st.warning("Data lampiran masih kosong!")
    else:
        st.info("Sedang diproses....")
        with st.spinner('Mengunduh template dan memproses dokumen...'):
            try:
                response = requests.get(GITHUB_RAW_URL)
                if response.status_code != 200:
                    st.error(f"Gagal mengunduh template (Status: {response.status_code}).")
                else:
                    template_file = io.BytesIO(response.content)
                    doc = Document(template_file)
                    
                    replacements = {
                        "{tentang}": tentang,
                        "{pelaksanaan}": pelaksanaan,
                        "{pelaksanan}": pelaksanaan, 
                        "{petugas}": petugas,
                        "{tanggal}": tanggal,
                        "{nomor}": nomor
                    }
                    
                    def replace_text_in_paragraphs(paragraphs):
                        for p in paragraphs:
                            original_text = p.text
                            changed = False
                            for key, value in replacements.items():
                                if key in original_text:
                                    original_text = original_text.replace(key, value)
                                    changed = True
                            if changed:
                                p.text = original_text
                                for run in p.runs:
                                    apply_bookman_font(run, size=12)

                    replace_text_in_paragraphs(doc.paragraphs)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if "{nama}" not in cell.text.lower():
                                    replace_text_in_paragraphs(cell.paragraphs)

                    target_table = None
                    template_row_idx = -1
                    for table in doc.tables:
                        for i, row in enumerate(table.rows):
                            for cell in row.cells:
                                if "{nama}" in cell.text.lower():
                                    target_table = table
                                    template_row_idx = i
                                    break
                            if target_table: break
                        if target_table: break
                    
                    if target_table and template_row_idx != -1:
                        for i in range(template_row_idx):
                            set_repeat_table_header(target_table.rows[i])

                        current_tr = target_table.rows[template_row_idx]._tr
                        total_data = len(edited_df)
                        
                        for index, row_data in edited_df.iterrows():
                            if index == 0:
                                target_row = target_table.rows[template_row_idx]
                            else:
                                spacer_row = target_table.add_row()
                                for cell in spacer_row.cells:
                                    p = cell.paragraphs[0]
                                    p.paragraph_format.space_before = Pt(0)
                                    p.paragraph_format.space_after = Pt(0)
                                    p.paragraph_format.line_spacing = 1.5 
                                    p.paragraph_format.keep_with_next = True 
                                    run = p.add_run("")
                                    apply_bookman_font(run, size=12)
                                
                                current_tr.addnext(spacer_row._tr)
                                current_tr = spacer_row._tr
                                
                                new_row = target_table.add_row()
                                current_tr.addnext(new_row._tr)
                                current_tr = new_row._tr 
                                target_row = new_row
                            
                            is_last_row = (index == total_data - 1)
                            set_cell_text_with_font(target_row.cells[0], f"{index + 1}.", keep_next=is_last_row)
                            set_cell_text_with_font(target_row.cells[1], str(row_data["Nama/Jabatan"]), keep_next=is_last_row)
                            set_cell_text_with_font(target_row.cells[2], str(row_data["NIP/Golongan"]), keep_next=is_last_row)
                            set_cell_text_with_font(target_row.cells[3], str(row_data["Posisi"]), keep_next=is_last_row)
                            set_cell_text_with_font(target_row.cells[4], f"Rp{row_data['Honor']}", keep_next=is_last_row)

                        final_spacer_row = target_table.add_row()
                        for cell in final_spacer_row.cells:
                            p = cell.paragraphs[0]
                            p.paragraph_format.line_spacing = 1 
                            p.paragraph_format.keep_with_next = True 
                            run = p.add_run("")
                            apply_bookman_font(run, size=12)
                        current_tr.addnext(final_spacer_row._tr)

                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)
                    st.session_state.doc_data = doc_io.getvalue()
                    st.session_state.doc_name = f"SK_{tentang.replace(' ', '_')}.docx"
                    st.session_state.doc_ready = True
                    
            except Exception as e:
                st.error(f"Terjadi kesalahan teknis: {e}")
                st.session_state.doc_ready = False

# 4. Menampilkan Tombol Unduh
if st.session_state.doc_ready:
    st.success("✅ Dokumen berhasil diproses")
    st.download_button(
        label="⬇️ Unduh Dokumen Hasil",
        data=st.session_state.doc_data,
        file_name=st.session_state.doc_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
