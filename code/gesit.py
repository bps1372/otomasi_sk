# 13 Mei 2026 13.44 Updated

import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH # Tambahan untuk rata tengah
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import requests
import copy

st.set_page_config(page_title="GESIT - BPS1372", layout="centered")

st.title("📝 GESIT (Generate SK Instan)")
st.write("Aplikasi ini akan mengotomasi pengisian dokumen SK Kuasa Pengguna Anggaran BPS Kota Solok.")

# --- KONFIGURASI GITHUB ---
GITHUB_RAW_URL = "https://raw.githubusercontent.com/bps1372/otomasi_sk/main/TemplateDokumen.docx"

# 1. Input Data Umum
st.subheader("Informasi Umum Dokumen")
col1, col2 = st.columns(2)

with col1:
    nomor = st.text_input("A.Nomor Dokumen", placeholder="Contoh: 042.1 TAHUN 2026")
    tanggal = st.text_input("B.Tanggal Ditetapkan", placeholder="Contoh: 15 Januari 2026")
    petugas = st.text_input("C.Menetapkan...", placeholder="Contoh: Petugas Sensus Ekonomi 2026")

with col2:
    tentang = st.text_input("D.Judul [tulis dengan huruf besar]", placeholder="Contoh: PETUGAS SENSUS EKONOMI 2026")
    pelaksanaan = st.text_input("E.Pelaksanaan Kegiatan", placeholder="Contoh: Sensus Ekonomi 2026")

# 2. Input Data Tabel (Dinamis dengan Fitur Tambah Baris & Kolom)
st.subheader("F. Lampiran")

# Inisialisasi Session State
if "df_lampiran" not in st.session_state:
    st.session_state.df_lampiran = pd.DataFrame({
        "Nama/Jabatan": ["Bapak WXYZABCDEFGHI Statistisi Ahli Muda"],
        "NIP/Golongan": ["200202192019031001 Gol: III/b"],
        "Posisi": ["PML"],
        "Honor": ["50.000"]
    })

# State untuk menyimpan data hasil editan
if "edited_df" not in st.session_state:
    st.session_state.edited_df = st.session_state.df_lampiran

# State khusus untuk Nama Kolom urut (pertama)
if "nama_kolom_no" not in st.session_state:
    st.session_state.nama_kolom_no = "No."

# --- PANEL PENGATURAN BARIS & KOLOM ---
st.markdown("##### Pengaturan Baris & Kolom")
col_n1, col_n2 = st.columns(2)

# Kolom Kiri: Tambah Baris
with col_n1:
    st.markdown("**1. Tambah Baris Kosong**")
    baris_baru = st.number_input("Jumlah baris kosong:", min_value=1, max_value=500, value=1)
    if st.button("Tambahkan Baris"):
        new_data_dict = {col: [""] * baris_baru for col in st.session_state.edited_df.columns}
        new_data = pd.DataFrame(new_data_dict)
        st.session_state.df_lampiran = pd.concat([st.session_state.edited_df, new_data], ignore_index=True)
        st.rerun()

# Kolom Kanan: Tambah & Edit SEMUA Kolom (Maks +4 Kolom Baru)
with col_n2:
    st.markdown("**2. Pengaturan Kolom (Maks. Tambah 4)**")
    current_cols = list(st.session_state.edited_df.columns)
    
    jumlah_kolom_tambahan = len(current_cols) - 4
    
    if jumlah_kolom_tambahan < 4:
        kolom_baru = st.text_input("Nama Kolom Baru:", key="input_tambah_kolom")
        if st.button("➕ Tambah Kolom"):
            if kolom_baru and kolom_baru not in current_cols:
                st.session_state.edited_df[kolom_baru] = ""
                st.session_state.df_lampiran = st.session_state.edited_df
                st.rerun()
    else:
        st.info("Batas maksimal 4 kolom tambahan telah tercapai.")

    st.markdown("---")
    
    # Gabungkan kolom urut (No) dengan kolom DataFrame ke dalam opsi dropdown
    pilihan_kolom = [st.session_state.nama_kolom_no] + current_cols
    kolom_lama = st.selectbox("Pilih kolom tabel untuk diubah namanya:", pilihan_kolom)
    nama_baru = st.text_input("Ubah nama menjadi:", key="input_ubah_kolom")
    
    if st.button("📝 Simpan Nama Baru"):
        if nama_baru and nama_baru not in pilihan_kolom:
            # Jika yang dipilih adalah kolom urut (pertama)
            if kolom_lama == st.session_state.nama_kolom_no:
                st.session_state.nama_kolom_no = nama_baru
            # Jika yang dipilih adalah kolom data lainnya
            else:
                st.session_state.edited_df = st.session_state.edited_df.rename(columns={kolom_lama: nama_baru})
                st.session_state.df_lampiran = st.session_state.edited_df
            st.rerun()
# -----------------------------------------------------

# Editor Tabel dengan num_rows="dynamic"
st.markdown("##### Preview Data Lampiran")
edited_df = st.data_editor(
    st.session_state.df_lampiran, 
    num_rows="dynamic", 
    use_container_width=True, 
    key="editor_key"
)

st.session_state.edited_df = edited_df

# --- STATE UNTUK DOKUMEN ---
if "doc_ready" not in st.session_state:
    st.session_state.doc_ready = False
if "doc_data" not in st.session_state:
    st.session_state.doc_data = None
if "doc_name" not in st.session_state:
    st.session_state.doc_name = ""

# --- FUNGSI CUSTOM FONT, XML, BORDER & ALIGNMENT ---
def apply_bookman_font(run, size=11, bold=False):
    run.font.name = 'Bookman Old Style'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Bookman Old Style')
    run.font.size = Pt(size)
    if bold:
        run.bold = True

def set_cell_text_with_font(cell, text, keep_next=False, bold=False, center=False):
    cell.text = text
    
    # Pengaturan perataan vertikal (tengah atas-bawah)
    if center:
        tcPr = cell._tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), "center")
        tcPr.append(vAlign)

    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        
        # Pengaturan perataan horizontal (tengah kiri-kanan)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        if keep_next:
            p.paragraph_format.keep_with_next = True
            
        for run in p.runs:
            apply_bookman_font(run, size=12, bold=bold)

def set_repeat_table_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)

def copy_cell_style(source_cell, target_cell):
    """Fungsi ajaib untuk menyalin garis (border) dari satu sel ke sel lainnya dengan aman"""
    source_tcPr = source_cell._tc.get_or_add_tcPr()
    target_tcPr = target_cell._tc.get_or_add_tcPr()
    
    for child in list(target_tcPr):
        target_tcPr.remove(child)
        
    for child in source_tcPr:
        target_tcPr.append(copy.deepcopy(child))
# -------------------------------------

# 3. Tombol Eksekusi
if st.button("Proses & Buat Dokumen", type="primary"):
    if edited_df.empty:
        st.warning("Data lampiran masih kosong!")
    else:
        st.info("Sedang diproses....")
        with st.spinner('Mengunduh template dan memproses dokumen...'):
            try:
                response = requests.get(GITHUB_RAW_URL)
                if response.status_code != 200:
                    st.error(f"Gagal mengunduh template (Status: {response.status_code}).")
                else:
                    template_file = io.BytesIO(response.content)
                    doc = Document(template_file)
                    
                    replacements = {
                        "{tentang}": tentang,
                        "{pelaksanaan}": pelaksanaan,
                        "{pelaksanan}": pelaksanaan, 
                        "{petugas}": petugas,
                        "{tanggal}": tanggal,
                        "{nomor}": nomor
                    }
                    
                    def replace_text_in_paragraphs(paragraphs):
                        for p in paragraphs:
                            original_text = p.text
                            changed = False
                            for key, value in replacements.items():
                                if key in original_text:
                                    original_text = original_text.replace(key, value)
                                    changed = True
                            if changed:
                                p.text = original_text
                                for run in p.runs:
                                    apply_bookman_font(run, size=12)

                    replace_text_in_paragraphs(doc.paragraphs)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if "{nama}" not in cell.text.lower():
                                    replace_text_in_paragraphs(cell.paragraphs)

                    target_table = None
                    template_row_idx = -1
                    for table in doc.tables:
                        for i, row in enumerate(table.rows):
                            for cell in row.cells:
                                if "{nama}" in cell.text.lower():
                                    target_table = table
                                    template_row_idx = i
                                    break
                            if target_table: break
                        if target_table: break
                    
                    if target_table and template_row_idx != -1:
                        for i in range(template_row_idx):
                            set_repeat_table_header(target_table.rows[i])

                        current_table_cols = len(target_table.columns)
                        needed_cols = len(edited_df.columns) + 1 # +1 untuk indeks Nomor (No.)

                        # 1. Tambah Kolom & Copy Garis (Border)
                        if needed_cols > current_table_cols:
                            cols_to_add = needed_cols - current_table_cols
                            for _ in range(cols_to_add):
                                target_table.add_column(Cm(2.5)) 
                            
                            for row in target_table.rows:
                                for new_col_idx in range(current_table_cols, len(row.cells)):
                                    source_cell = row.cells[current_table_cols - 1]
                                    target_cell = row.cells[new_col_idx]
                                    copy_cell_style(source_cell, target_cell)
                            
                        # 2. Tulis Ulang Header (DENGAN BOLD & RATA TENGAH)
                        if len(target_table.rows) > 0:
                            header_text_row = target_table.rows[0]
                            
                            # UPDATE KHUSUS UNTUK KOLOM NO:
                            if len(header_text_row.cells) > 0:
                                set_cell_text_with_font(header_text_row.cells[0], st.session_state.nama_kolom_no, keep_next=True, bold=True, center=True)
                            
                            for col_idx, col_name in enumerate(edited_df.columns):
                                cell_idx = col_idx + 1
                                if cell_idx < len(header_text_row.cells):
                                    # Tambahan bold=True dan center=True khusus Header
                                    set_cell_text_with_font(header_text_row.cells[cell_idx], col_name, keep_next=True, bold=True, center=True)
                                    
                        if len(target_table.rows) > 1:
                            header_num_row = target_table.rows[1]
                            if len(header_num_row.cells) > 1 and "(" in header_num_row.cells[1].text and ")" in header_num_row.cells[1].text:
                                for col_idx in range(current_table_cols - 1, needed_cols - 1):
                                    cell_idx = col_idx + 1
                                    if cell_idx < len(header_num_row.cells):
                                        # Penomoran baris header juga dibuat rata tengah
                                        set_cell_text_with_font(header_num_row.cells[cell_idx], f"({cell_idx + 1})", keep_next=True, center=True)
                        # ----------------------------------------------

                        current_tr = target_table.rows[template_row_idx]._tr
                        total_data = len(edited_df)
                        
                        for index, row_data in edited_df.iterrows():
                            if index == 0:
                                target_row = target_table.rows[template_row_idx]
                            else:
                                spacer_row = target_table.add_row()
                                for cell in spacer_row.cells:
                                    p = cell.paragraphs[0]
                                    p.paragraph_format.space_before = Pt(0)
                                    p.paragraph_format.space_after = Pt(0)
                                    p.paragraph_format.line_spacing = 1.5 
                                    p.paragraph_format.keep_with_next = True 
                                    run = p.add_run("")
                                    apply_bookman_font(run, size=12)
                                
                                current_tr.addnext(spacer_row._tr)
                                current_tr = spacer_row._tr
                                
                                new_row = target_table.add_row()
                                current_tr.addnext(new_row._tr)
                                current_tr = new_row._tr 
                                target_row = new_row
                            
                            is_last_row = (index == total_data - 1)
                            
                            # Baris Data Tetap Normal (Tanpa Bold / Rata Tengah, kecuali jika diinginkan)
                            set_cell_text_with_font(target_row.cells[0], f"{index + 1}.", keep_next=is_last_row)
                            
                            for col_idx, col_name in enumerate(edited_df.columns):
                                val = str(row_data[col_name])
                                if ("honor" in col_name.lower() or "uang" in col_name.lower() or "tarif" in col_name.lower()) and val and not val.lower().startswith("rp"):
                                    val = f"Rp{val}"
                                set_cell_text_with_font(target_row.cells[col_idx + 1], val, keep_next=is_last_row)

                        final_spacer_row = target_table.add_row()
                        for cell in final_spacer_row.cells:
                            p = cell.paragraphs[0]
                            p.paragraph_format.line_spacing = 1 
                            p.paragraph_format.keep_with_next = True 
                            run = p.add_run("")
                            apply_bookman_font(run, size=12)
                        current_tr.addnext(final_spacer_row._tr)

                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)
                    st.session_state.doc_data = doc_io.getvalue()
                    st.session_state.doc_name = f"SK_{tentang.replace(' ', '_')}.docx"
                    st.session_state.doc_ready = True
                    
            except Exception as e:
                st.error(f"Terjadi kesalahan teknis: {e}")
                st.session_state.doc_ready = False

# 4. Menampilkan Tombol Unduh
if st.session_state.doc_ready:
    st.success("✅ Dokumen berhasil diproses")
    st.download_button(
        label="⬇️ Unduh Dokumen Hasil",
        data=st.session_state.doc_data,
        file_name=st.session_state.doc_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
st.write("")
st.write("")
st.write("")
st.write("")

st.write("                       Copyright @BPS Kota Solok")
